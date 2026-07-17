"""
Document reader tool for reading images and documents from disk.

Supported formats:
- Images: PNG, JPG, GIF, WebP, BMP, TIFF
- Documents: PDF, Word (.docx), Excel (.xlsx, .xls), HTML (.html, .htm)
- PDF: text extraction, page rendering, figure/table extraction by caption
- Word/Excel: text, tables, embedded images
- HTML: rendered in chat
"""

import zipfile
from pathlib import Path
from typing import Any, Dict, List

# Supported image extensions
IMAGE_EXTENSIONS = frozenset({
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif",
})

# Supported document extensions
PDF_EXTENSIONS = frozenset({".pdf"})
WORD_EXTENSIONS = frozenset({".docx"})
EXCEL_XLSX_EXTENSIONS = frozenset({".xlsx"})
EXCEL_XLS_EXTENSIONS = frozenset({".xls"})
HTML_EXTENSIONS = frozenset({".html", ".htm"})

EXT_TO_MIME = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".tiff": "image/tiff",
    ".tif": "image/tiff",
}

DOCUMENT_TOOL_NAMES = frozenset({"read_document"})


def get_tool_definitions() -> List[Dict[str, Any]]:
    """Tool definitions for document reading."""
    return [
        {
            "name": "read_document",
            "description": "Read documents from disk. Supports: images, PDF, Word, Excel, HTML. Infer params from user intent—users ask countless ways. Full doc: omit params. Specific part: PDF page=N, figure_number=N, section='X'; Word table_name='X', section='X'; Excel sheet_name='X'. Path relative or absolute.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filepath": {
                        "type": "string",
                        "description": "Path to file (e.g. paper/logo.png, image/new_paper.pdf, or absolute path)",
                    },
                    "page": {
                        "type": "integer",
                        "description": "PDF: show only this page (1-indexed). Infer from 'page 5', 'page 10', 'fifth page', etc. Omit for full document.",
                    },
                    "figure_number": {
                        "type": "integer",
                        "description": "PDF: show only the Nth figure/image/table/graph (1-indexed). Infer from 'figure 5', 'fig 3', 'graph 1', 'table 2', 'image 4', 'tab 1', etc. Omit for full document.",
                    },
                    "table_name": {
                        "type": "string",
                        "description": "Word (.docx): show only the table whose caption, header, or content contains this text. Infer from 'table about X', 'Characteristics table', etc. Omit for full document.",
                    },
                    "section": {
                        "type": "string",
                        "description": "PDF/Word: show only the text section containing this string. Infer from 'methodology section', 'subsection on results', 'part about X', etc. Omit for full document.",
                    },
                    "sheet_name": {
                        "type": "string",
                        "description": "Excel: show only this sheet. Infer from 'sheet Sales', 'Q4 data', sheet name, etc. Omit for all sheets.",
                    },
                },
                "required": ["filepath"],
            },
        },
    ]


def _resolve_filepath(filepath: str, project_root: Path) -> Path:
    """Resolve document paths, including app-root-style paths such as /paper/x.pdf."""
    raw = str(filepath or "").strip()
    p = Path(raw).expanduser()
    if not p.is_absolute():
        return (project_root / p).resolve()

    absolute = p.resolve()
    if absolute.exists():
        return absolute

    # In the Autonomous Lab, users often write '/paper/paper.pdf' to mean a
    # path below the application root rather than the host filesystem root.
    project_candidate = (project_root / raw.lstrip("/\\")).resolve()
    if project_candidate.exists():
        return project_candidate
    return absolute


def _read_image(p: Path) -> Dict[str, Any]:
    """Read image file and return artifact-ready dict."""
    if not p.exists():
        return {"error": f"File not found: {p}"}
    if not p.is_file():
        return {"error": f"Not a file: {p}"}

    ext = p.suffix.lower()
    mime_type = EXT_TO_MIME.get(ext, "image/png")

    try:
        image_bytes = p.read_bytes()
    except OSError as e:
        return {"error": f"Cannot read file: {e}"}

    if len(image_bytes) == 0:
        return {"error": "File is empty"}

    # Maximum file size to prevent excessive memory usage
    max_bytes = 10 * 1024 * 1024
    if len(image_bytes) > max_bytes:
        return {"error": f"Image too large ({len(image_bytes) / 1024 / 1024:.1f} MB). Max 10 MB."}

    return {
        "artifact_type": "image_file",
        "figure_image": {"mime_type": mime_type, "data": image_bytes},
        "artifact_title": f"Image: {p.name}",
        "message": f"Read image from {p.name}",
    }


def _extract_section_from_text(text: str, query: str, context_chars: int = 1500) -> str | None:
    """Find and return the text chunk containing query, with surrounding context."""
    if not text or not query or not query.strip():
        return None
    q = query.strip().lower()
    t_lower = text.lower()
    idx = t_lower.find(q)
    if idx < 0:
        return None
    start = max(0, idx - context_chars)
    end = min(len(text), idx + len(query) + context_chars)
    chunk = text[start:end]
    if start > 0:
        chunk = "..." + chunk
    if end < len(text):
        chunk = chunk + "..."
    return chunk


def _read_pdf(p: Path, page: int | None = None, figure_number: int | None = None, section: str | None = None) -> Dict[str, Any]:
    """Read PDF: extract text AND render pages or a specific figure. page/figure_number limit output."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {"error": "PyMuPDF not installed. Run: pip install PyMuPDF"}

    if not p.exists() or not p.is_file():
        return {"error": f"File not found or not a file: {p}"}

    max_text_pages = 50
    max_image_pages = 20
    max_chars = 100_000
    zoom = 2.0
    zoom_figure = 2.5  # Higher resolution for single-figure extraction

    try:
        doc = fitz.open(p)
        page_count = len(doc)

        # Locate page containing the figure/table caption (e.g., "Figure 1:", "Table 2.")
        if figure_number is not None and figure_number >= 1:
            prefixes = ("figure", "fig.", "fig", "graph", "table", "tab.", "tab", "image", "img.")
            caption_patterns = [f"{p} {figure_number}:" for p in prefixes] + [f"{p} {figure_number}." for p in prefixes]
            found_page_idx = None
            for i in range(page_count):
                t = doc[i].get_text().lower()
                if any(p in t for p in caption_patterns):
                    found_page_idx = i
                    break
            if found_page_idx is None:
                for i in range(page_count):
                    t = doc[i].get_text().lower()
                    if any(f"{p} {figure_number}" in t for p in prefixes):
                        found_page_idx = i
                        break
            page_idx = found_page_idx if found_page_idx is not None else min(figure_number, page_count) - 1
            page_idx = max(0, min(page_idx, page_count - 1))
            pg = doc[page_idx]

            # Extract largest embedded image on page (excludes small decorative elements)
            best_img: Dict[str, Any] | None = None
            best_area = 80 * 80  # Minimum pixel area threshold
            for img in pg.get_images():
                xref = img[0]
                try:
                    info = doc.extract_image(xref)
                    w = info.get("width", 0)
                    h = info.get("height", 0)
                    pix_bytes = info.get("image")
                    if not pix_bytes or w * h < best_area:
                        continue
                    ext = info.get("ext", "png")
                    mime = {"png": "image/png", "jpeg": "image/jpeg", "jpg": "image/jpeg", "gif": "image/gif"}.get(ext.lower(), "image/png")
                    if len(pix_bytes) < 10 * 1024 * 1024 and w * h > best_area:
                        best_area = w * h
                        best_img = {"mime_type": mime, "data": pix_bytes}
                except Exception:
                    continue

            if best_img:
                doc.close()
                return {
                    "artifact_type": "pdf_document",
                    "artifact_content": f"Figure {figure_number} from {p.name}",
                    "artifact_title": f"PDF: {p.name} — Figure {figure_number}",
                    "figure_images": [best_img],
                    "message": f"Figure {figure_number} from {p.name}",
                }

            # Vector graphics: render full page (embedded image extraction not applicable)
            z = zoom_figure
            mat = fitz.Matrix(z, z)
            pix = pg.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes(output="png")
            doc.close()
            return {
                "artifact_type": "pdf_document",
                "artifact_content": f"Figure {figure_number} (page {page_idx + 1}) — vector graphic, full page shown",
                "artifact_title": f"PDF: {p.name} — Figure {figure_number}",
                "figure_images": [{"mime_type": "image/png", "data": img_bytes}],
                "message": f"Figure {figure_number} from {p.name} (vector figure, page shown)",
            }

        # Single-page extraction (1-indexed)
        if page is not None and page >= 1:
            page_idx = page - 1
            if page_idx >= page_count:
                doc.close()
                return {"error": f"PDF has {page_count} pages; page {page} does not exist."}
            pg = doc[page_idx]
            mat = fitz.Matrix(zoom, zoom)
            pix = pg.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes(output="png")
            text = pg.get_text()
            doc.close()
            return {
                "artifact_type": "pdf_document",
                "artifact_content": text[:max_chars] if text.strip() else f"Page {page} of {p.name}",
                "artifact_title": f"PDF: {p.name} — Page {page}",
                "figure_images": [{"mime_type": "image/png", "data": img_bytes}],
                "message": f"Page {page} of {p.name}",
            }

        # Full document: text extraction and page rendering
        text_parts = []
        total_chars = 0
        for i in range(min(page_count, max_text_pages)):
            if total_chars >= max_chars:
                text_parts.append(f"\n... [truncated after {max_chars} chars]")
                break
            pg = doc[i]
            block = pg.get_text()
            text_parts.append(block)
            total_chars += len(block)
        text = "\n\n".join(text_parts)

        if section and section.strip():
            chunk = _extract_section_from_text(text, section.strip())
            if chunk:
                doc.close()
                return {
                    "artifact_type": "pdf_document",
                    "artifact_content": chunk,
                    "artifact_title": f"PDF: {p.name} — Section: {section[:40]}",
                    "figure_images": [],
                    "message": f"Extracted section '{section[:40]}...' from {p.name}",
                }
            # Section not found; proceed with full document

        figure_images = []
        for i in range(min(page_count, max_image_pages)):
            pg = doc[i]
            mat = fitz.Matrix(zoom, zoom)
            pix = pg.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes(output="png")
            figure_images.append({"mime_type": "image/png", "data": img_bytes})
        doc.close()

        return {
            "artifact_type": "pdf_document",
            "artifact_content": text[:max_chars] if text.strip() else "(No extractable text; see page images below)",
            "artifact_title": f"PDF: {p.name}",
            "figure_images": figure_images,
            "message": f"Read PDF from {p.name} ({min(page_count, max_image_pages)} pages, text + figures)",
        }
    except Exception as e:
        return {"error": str(e)}


def _read_word(p: Path, table_name: str | None = None, section: str | None = None) -> Dict[str, Any]:
    """Read Word .docx: extract text and embedded images. table_name filters to one table."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    if not p.exists() or not p.is_file():
        return {"error": f"File not found or not a file: {p}"}

    max_chars = 100_000
    ext_to_mime = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                   ".gif": "image/gif", ".emf": "image/x-emf", ".wmf": "image/x-wmf"}

    try:
        doc = Document(p)

        # Extract table by caption or header match
        if table_name and table_name.strip():
            search = table_name.strip().lower()
            tbl_count = 0
            prev_para = ""
            for child in doc.element.body:
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag == "p":
                    prev_para = "".join(t.text or "" for t in child.iter() if hasattr(t, "text")).strip()
                elif tag == "tbl":
                    if tbl_count < len(doc.tables):
                        tbl = doc.tables[tbl_count]
                        header_text = " ".join(cell.text.strip() for row in tbl.rows[:1] for cell in row.cells).lower()
                        all_cells = " ".join(cell.text.strip() for row in tbl.rows for cell in row.cells).lower()
                        if search in prev_para.lower() or search in header_text or search in all_cells:
                            rows = []
                            for row in tbl.rows:
                                rows.append("| " + " | ".join(cell.text.strip() for cell in row.cells) + " |")
                            table_md = "\n".join(rows)
                            return {
                                "artifact_type": "word_document",
                                "artifact_content": table_md,
                                "artifact_title": f"Word: {p.name} — Table: {table_name[:50]}",
                                "figure_images": [],
                                "message": f"Extracted table '{table_name[:40]}...' from {p.name}",
                            }
                    tbl_count += 1
                    prev_para = ""
            return {"error": f"No table found matching '{table_name}'. Try a shorter or different search string."}

        # Extract text section by query match
        if section and section.strip():
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text.strip() for cell in row.cells))
            text = "\n\n".join(parts)
            chunk = _extract_section_from_text(text, section.strip())
            if chunk:
                return {
                    "artifact_type": "word_document",
                    "artifact_content": chunk,
                    "artifact_title": f"Word: {p.name} — Section: {section[:40]}",
                    "figure_images": [],
                    "message": f"Extracted section '{section[:40]}...' from {p.name}",
                }

        # Full document extraction
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n\n".join(parts)

        # Extract embedded images from OOXML package (word/media/)
        figure_images = []
        try:
            with zipfile.ZipFile(p, "r") as z:
                for name in z.namelist():
                    if not name.startswith("word/media/") or name.endswith("/"):
                        continue
                    data = z.read(name)
                    ext = Path(name).suffix.lower()
                    mime = ext_to_mime.get(ext, "image/png")
                    if len(data) < 10 * 1024 * 1024:  # 10 MB limit per image
                        figure_images.append({"mime_type": mime, "data": data})
        except Exception:
            pass

        return {
            "artifact_type": "word_document",
            "artifact_content": text[:max_chars] if text.strip() else "(Document empty; see embedded images below)",
            "artifact_title": f"Word: {p.name}",
            "figure_images": figure_images,
            "message": f"Read Word from {p.name} (text + {len(figure_images)} image(s))",
        }
    except Exception as e:
        return {"error": str(e)}


def _read_excel(p: Path, sheet_name: str | None = None) -> Dict[str, Any]:
    """Read Excel .xlsx and extract tables as markdown. sheet_name filters to one sheet."""
    try:
        import openpyxl
    except ImportError:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}

    if not p.exists() or not p.is_file():
        return {"error": f"File not found or not a file: {p}"}

    max_sheets = 5
    max_rows_per_sheet = 500
    max_chars = 100_000

    try:
        wb = openpyxl.load_workbook(p, read_only=False, data_only=True)
        all_sheets = wb.sheetnames
        if sheet_name and sheet_name.strip():
            q = sheet_name.strip().lower()
            matches = [s for s in all_sheets if q in s.lower()]
            sheets = matches[:1] if matches else all_sheets[:1]
            if not matches:
                wb.close()
                return {"error": f"No sheet matching '{sheet_name}'. Available: {', '.join(all_sheets[:10])}"}
        else:
            sheets = all_sheets[:max_sheets]
        parts = []
        total_chars = 0
        for sheet_name in sheets:
            if total_chars >= max_chars:
                break
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            parts.append(f"### Sheet: {sheet_name}")
            header = rows[0]
            header_str = "| " + " | ".join(str(h) if h is not None else "" for h in header) + " |"
            parts.append(header_str)
            parts.append("| " + " | ".join("---" for _ in header) + " |")
            for row in rows[1:max_rows_per_sheet]:
                row_str = "| " + " | ".join(str(c) if c is not None else "" for c in row) + " |"
                parts.append(row_str)
                total_chars += len(row_str)
                if total_chars >= max_chars:
                    parts.append("... (truncated)")
                    break
            parts.append("")
        wb.close()
        text = "\n".join(parts)

        # Extract embedded images from OOXML package (xl/media/)
        figure_images = []
        try:
            with zipfile.ZipFile(p, "r") as z:
                for name in z.namelist():
                    if not name.startswith("xl/media/") or name.endswith("/"):
                        continue
                    data = z.read(name)
                    ext = Path(name).suffix.lower()
                    mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                            "gif": "image/gif", "emf": "image/x-emf", "wmf": "image/x-wmf"}.get(ext[1:] if ext else "", "image/png")
                    if len(data) < 10 * 1024 * 1024:
                        figure_images.append({"mime_type": mime, "data": data})
        except Exception:
            pass

        return {
            "artifact_type": "excel_document",
            "artifact_content": text[:max_chars] if text.strip() else "(Sheets empty; see embedded images below)",
            "artifact_title": f"Excel: {p.name}",
            "figure_images": figure_images,
            "message": f"Read Excel from {p.name} ({len(sheets)} sheet(s), {len(figure_images)} image(s))",
        }
    except Exception as e:
        return {"error": str(e)}


def _read_excel_xls(p: Path, sheet_name: str | None = None) -> Dict[str, Any]:
    """Read legacy Excel .xls via xlrd2. sheet_name filters to one sheet."""
    try:
        import xlrd
    except ImportError:
        return {"error": "xlrd2 not installed. Run: pip install xlrd2"}

    if not p.exists() or not p.is_file():
        return {"error": f"File not found or not a file: {p}"}

    max_sheets = 5
    max_rows_per_sheet = 500
    max_chars = 100_000

    try:
        book = xlrd.open_workbook(str(p))
        all_sheets = book.sheet_names()
        if sheet_name and sheet_name.strip():
            q = sheet_name.strip().lower()
            matches = [s for s in all_sheets if q in s.lower()]
            sheet_names = matches[:1] if matches else []
            if not matches:
                return {"error": f"No sheet matching '{sheet_name}'. Available: {', '.join(all_sheets[:10])}"}
        else:
            sheet_names = all_sheets[:max_sheets]
        parts = []
        total_chars = 0
        for sheet_name in sheet_names:
            if total_chars >= max_chars:
                break
            sh = book.sheet_by_name(sheet_name)
            rows = []
            for row_idx in range(min(sh.nrows, max_rows_per_sheet)):
                row = sh.row_values(row_idx)
                rows.append(row)
            if not rows:
                continue
            parts.append(f"### Sheet: {sheet_name}")
            header = rows[0]
            n_cols = max(len(r) for r in rows) if rows else 0
            header = [str(header[i]) if i < len(header) else "" for i in range(n_cols)]
            header_str = "| " + " | ".join(str(h) if h else "" for h in header) + " |"
            parts.append(header_str)
            parts.append("| " + " | ".join("---" for _ in header) + " |")
            for row in rows[1:]:
                row_padded = [str(row[i]) if i < len(row) else "" for i in range(n_cols)]
                row_str = "| " + " | ".join(str(c) if c else "" for c in row_padded) + " |"
                parts.append(row_str)
                total_chars += len(row_str)
                if total_chars >= max_chars:
                    parts.append("... (truncated)")
                    break
            parts.append("")
        text = "\n".join(parts)
        if not text.strip():
            return {"error": "Excel file appears empty"}
        return {
            "artifact_type": "excel_document",
            "artifact_content": text[:max_chars] if text.strip() else "(Sheets empty)",
            "artifact_title": f"Excel: {p.name}",
            "figure_images": [],  # Legacy .xls format does not support embedded images
            "message": f"Read Excel (.xls) from {p.name} ({len(sheet_names)} sheet(s))",
        }
    except Exception as e:
        return {"error": str(e)}


def _read_html(p: Path) -> Dict[str, Any]:
    """Read HTML file and return as report_html for rendered display in chat."""
    if not p.exists() or not p.is_file():
        return {"error": f"File not found or not a file: {p}"}

    max_chars = 2_000_000  # Maximum content size for HTML (e.g., Plotly exports)

    try:
        html_content = p.read_text(encoding="utf-8", errors="replace")
        if not html_content.strip():
            return {"error": "HTML file is empty"}
        return {
            "artifact_type": "report_html",
            "artifact_content": html_content[:max_chars],
            "artifact_title": f"HTML: {p.name}",
            "message": f"Read HTML from {p.name} (rendered below)",
        }
    except Exception as e:
        return {"error": str(e)}


def execute_tool(
    name: str,
    args: Dict[str, Any],
    project_root: Path,
    session_context: Any = None,
) -> Any:
    """Execute document tool. Returns artifact dict or error string."""
    if name != "read_document":
        return f"Error: Unknown tool '{name}'"

    filepath = args.get("filepath", "")
    if not filepath or not str(filepath).strip():
        return "Error: filepath is required"

    p = _resolve_filepath(str(filepath).strip(), project_root)

    if not p.exists():
        return f"Error: File not found: {filepath}"
    if not p.is_file():
        return f"Error: Not a file: {filepath}"

    ext = p.suffix.lower()

    if ext in IMAGE_EXTENSIONS:
        result = _read_image(p)
        if "error" in result:
            return f"Error: {result['error']}"
        return result

    if ext in PDF_EXTENSIONS:
        page = args.get("page")
        figure_number = args.get("figure_number")
        section = (args.get("section") or "").strip() or None
        if page is not None:
            try:
                page = int(page)
            except (TypeError, ValueError):
                page = None
        if figure_number is not None:
            try:
                figure_number = int(figure_number)
            except (TypeError, ValueError):
                figure_number = None
        result = _read_pdf(p, page=page, figure_number=figure_number, section=section)
        if "error" in result:
            return f"Error: {result['error']}"
        return result

    if ext in WORD_EXTENSIONS:
        table_name = (args.get("table_name") or "").strip() or None
        section = (args.get("section") or "").strip() or None
        result = _read_word(p, table_name=table_name, section=section)
        if "error" in result:
            return f"Error: {result['error']}"
        return result

    if ext in EXCEL_XLSX_EXTENSIONS:
        sheet_name = (args.get("sheet_name") or "").strip() or None
        result = _read_excel(p, sheet_name=sheet_name)
        if "error" in result:
            return f"Error: {result['error']}"
        return result

    if ext in EXCEL_XLS_EXTENSIONS:
        sheet_name = (args.get("sheet_name") or "").strip() or None
        result = _read_excel_xls(p, sheet_name=sheet_name)
        if "error" in result:
            return f"Error: {result['error']}"
        return result

    if ext in HTML_EXTENSIONS:
        result = _read_html(p)
        if "error" in result:
            return f"Error: {result['error']}"
        return result

    supported = sorted(IMAGE_EXTENSIONS | PDF_EXTENSIONS | WORD_EXTENSIONS | EXCEL_XLSX_EXTENSIONS | EXCEL_XLS_EXTENSIONS | HTML_EXTENSIONS)
    return f"Error: Unsupported file type '{ext}'. Supported: {', '.join(supported)}"
